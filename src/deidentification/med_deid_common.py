# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import json
import logging
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import numpy as np
from PIL import Image
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

GLAUCOMA_TYPES = ("AACG", "CACG", "POAG")

DOCTOR_TRIGGER_LABELS = [
    "住院医生", "住院医师", "主治医师", "主任医师", "副主任医师",
    "上级医师", "经治医师", "接诊医师", "值班医师", "申请医生",
    "报告医生", "审核医师", "手术者", "手术医师", "医师签名",
    "签名", "审签者", "记录者", "报告者",
]

DOCTOR_TITLE_WORDS = [
    "主任医师", "副主任医师", "主治医师", "住院医师", "住院医生", "医师", "医生"
]

REMOVE_VALUE_LABELS = [
    "床号", "病案号", "门诊号", "门诊病案号", "住院号", "病历号",
    "联系电话", "电话预约", "联系方式", "联系人电话", "住址", "现住址",
    "出生地点", "出生地", "户口地址", "家庭住址", "单位地址",
    "联系人地址", "医院地址", "病史陈述者",
]

PATIENT_NAME_LABELS = ["姓名", "患者姓名"]
HOSPITAL_LABELS = ["医院名称", "就诊医院", "医院地址", "医院", "公众号", "电子病历系统"]

ALL_KNOWN_LABELS = sorted(
    set(DOCTOR_TRIGGER_LABELS) | set(REMOVE_VALUE_LABELS) | set(PATIENT_NAME_LABELS) |
    set(HOSPITAL_LABELS) | {"性别", "年龄", "科室", "婚姻", "职业", "民族", "出生日期", "出生地点",
                            "入院时间", "记录时间", "与患者关系", "是否可靠", "诊断", "备注", "门诊日期", "就诊日期"},
    key=len,
    reverse=True,
)

A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

BLIP_XPATH = f".//{{{A_NS}}}blip"
DOC_REL_XPATH = f".//{{{PKG_REL_NS}}}Relationship"


@dataclass
class Block:
    location: str
    kind: str
    text: str
    rows: Optional[List[List[str]]] = None


def setup_logger(log_path: Path) -> logging.Logger:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger(str(log_path))
    logger.setLevel(logging.INFO)
    logger.propagate = False
    if logger.handlers:
        for h in list(logger.handlers):
            logger.removeHandler(h)
    fmt = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s", "%Y-%m-%d %H:%M:%S")
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setFormatter(fmt)
    sh = logging.StreamHandler()
    sh.setFormatter(fmt)
    logger.addHandler(fh)
    logger.addHandler(sh)
    return logger


def truthy(v: str) -> bool:
    return str(v).strip().lower() in {"1", "true", "yes", "y", "是"}


def normalize_spaces(text: str) -> str:
    if text is None:
        return ""
    text = text.replace("\u3000", " ").replace("\xa0", " ").replace("\ufeff", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[ ]*\n[ ]*", "\n", text)
    return text.strip()


def normalize_spaces_keep_tabs(text: str) -> str:
    """规范空白但保留制表符，避免表格列结构被抹平。"""
    if text is None:
        return ""
    text = text.replace("\u3000", " ").replace("\xa0", " ").replace("\ufeff", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ ]+", " ", text)
    text = re.sub(r"[ ]*\n[ ]*", "\n", text)
    text = re.sub(r"\t+", "\t", text)
    text = re.sub(r"[ ]*\t[ ]*", "\t", text)
    return text.strip()


def compact_text(text: str) -> str:
    text = normalize_spaces(text)
    return re.sub(r"[\s:：\-—_/\\|·•，,。；;（）()\[\]【】]+", "", text)


def normalize_person_name(name: str) -> str:
    if not name:
        return ""
    name = normalize_spaces(name)
    name = re.sub(r"(主任医师|副主任医师|主治医师|住院医师|住院医生|医师|医生|教授|门诊|专科|签名)$", "", name)
    name = re.sub(r"^[：:\-—\s]+|[：:\-—\s]+$", "", name)
    m = re.search(r"[\u4e00-\u9fff]{2,4}", name)
    return m.group(0) if m else ""


def normalize_hospital_name(name: str) -> str:
    name = normalize_spaces(name)
    name = re.sub(r"[“”\"'<>《》\[\]【】]", "", name)
    return re.sub(r"\s+", "", name)


def safe_read_csv(path: Path) -> List[dict]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: List[dict], fieldnames: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            out = {k: row.get(k, "") for k in fieldnames}
            writer.writerow(out)


def discover_patients(input_root: Path) -> List[dict]:
    patients = []
    for glaucoma_type in GLAUCOMA_TYPES:
        type_dir = input_root / glaucoma_type
        if not type_dir.is_dir():
            continue
        for patient_dir in sorted([p for p in type_dir.iterdir() if p.is_dir()], key=lambda p: p.name):
            patients.append({
                "glaucoma_type": glaucoma_type,
                "patient_name": patient_dir.name,
                "patient_dir": patient_dir,
            })
    patients.sort(key=lambda x: (x["glaucoma_type"], x["patient_name"]))
    for idx, item in enumerate(patients):
        item["patient_id"] = f"PT_{idx:06d}"
    return patients


def infer_doc_type(path: Path) -> str:
    stem = path.stem
    if "入院" in stem:
        return "admission"
    if "出院" in stem:
        return "discharge"
    if "术后" in stem:
        return "followup_image_docx"
    return "unknown"


def iter_block_items(parent):
    from docx.document import Document as _Document
    from docx.section import _Header, _Footer
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, (_Header, _Footer)):
        parent_elm = parent._element
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def rows_to_text(rows: List[List[str]], escape_cell_newlines: bool = False) -> str:
    lines = []
    for row in rows:
        out_cells = []
        for cell in row:
            c = cell if cell is not None else ""
            if escape_cell_newlines:
                c = c.replace("\n", r"\n")
            out_cells.append(c)
        lines.append("\t".join(out_cells))
    return "\n".join(lines)


def _cell_text_with_nested_tables(cell, depth: int = 0, max_depth: int = 8) -> str:
    # python-docx 的 cell.text 不会包含子表格内容，这里递归展开。
    if depth >= max_depth:
        return normalize_spaces(cell.text)

    parts: List[str] = []
    for item in iter_block_items(cell):
        if isinstance(item, Paragraph):
            txt = normalize_spaces(item.text)
            if txt:
                parts.append(txt)
        elif isinstance(item, Table):
            nested_rows = table_to_rows(item, depth + 1, max_depth=max_depth)
            nested_text = rows_to_text(nested_rows, escape_cell_newlines=True)
            if nested_text:
                parts.append("<<NESTED_TABLE>>\n" + nested_text + "\n<</NESTED_TABLE>>")

    if parts:
        return normalize_spaces_keep_tabs("\n".join(parts))
    return normalize_spaces(cell.text)


def table_to_rows(table: Table, depth: int = 0, max_depth: int = 8) -> List[List[str]]:
    rows = []
    for row in table.rows:
        rows.append([_cell_text_with_nested_tables(cell, depth=depth, max_depth=max_depth) for cell in row.cells])
    return rows


def document_to_blocks(docx_path: Path) -> dict:
    result = {"header": [], "body": [], "footer": []}

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        logging.warning(f"跳过无法读取的文档 (可能为空或已损坏): {docx_path} | 错误信息: {e}")
        return result

    for idx, item in enumerate(iter_block_items(doc)):
        if isinstance(item, Paragraph):
            result["body"].append(
                Block(location=f"body.paragraph[{idx}]", kind="paragraph", text=normalize_spaces(item.text)))
        else:
            rows = table_to_rows(item)
            text = rows_to_text(rows)
            result["body"].append(Block(location=f"body.table[{idx}]", kind="table", text=text, rows=rows))

    for sec_idx, section in enumerate(doc.sections):
        for idx, item in enumerate(iter_block_items(section.header)):
            if isinstance(item, Paragraph):
                result["header"].append(Block(location=f"header[{sec_idx}].paragraph[{idx}]", kind="paragraph",
                                              text=normalize_spaces(item.text)))
            else:
                rows = table_to_rows(item)
                text = rows_to_text(rows)
                result["header"].append(
                    Block(location=f"header[{sec_idx}].table[{idx}]", kind="table", text=text, rows=rows))
        for idx, item in enumerate(iter_block_items(section.footer)):
            if isinstance(item, Paragraph):
                result["footer"].append(Block(location=f"footer[{sec_idx}].paragraph[{idx}]", kind="paragraph",
                                              text=normalize_spaces(item.text)))
            else:
                rows = table_to_rows(item)
                text = rows_to_text(rows)
                result["footer"].append(
                    Block(location=f"footer[{sec_idx}].table[{idx}]", kind="table", text=text, rows=rows))
    return result


def write_blocks_to_docx(structure: dict, output_path: Path, title: Optional[str] = None) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    if structure.get("header"):
        sec = doc.sections[0]
        for block in structure["header"]:
            if block.kind == "paragraph":
                sec.header.add_paragraph(block.text)
            else:
                rows = block.rows or []
                if rows:
                    tbl = sec.header.add_table(rows=len(rows), cols=len(rows[0]))
                    for r, row in enumerate(rows):
                        for c, cell_text in enumerate(row):
                            tbl.cell(r, c).text = cell_text
    for block in structure.get("body", []):
        if block.kind == "paragraph":
            doc.add_paragraph(block.text)
        else:
            rows = block.rows or []
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        tbl.cell(r, c).text = cell_text
    if structure.get("footer"):
        sec = doc.sections[0]
        for block in structure["footer"]:
            if block.kind == "paragraph":
                sec.footer.add_paragraph(block.text)
            else:
                rows = block.rows or []
                if rows:
                    tbl = sec.footer.add_table(rows=len(rows), cols=len(rows[0]))
                    for r, row in enumerate(rows):
                        for c, cell_text in enumerate(row):
                            tbl.cell(r, c).text = cell_text
    doc.save(str(output_path))


def write_plain_text_docx(text: str, output_path: Path, title: Optional[str] = None,
                          metadata: Optional[List[str]] = None) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    if metadata:
        for line in metadata:
            if line:
                doc.add_paragraph(line)
    for para in normalize_spaces(text).split("\n"):
        doc.add_paragraph(para)
    doc.save(str(output_path))


def extract_docx_images_in_order(docx_path: Path) -> List[Tuple[str, bytes]]:
    out: List[Tuple[str, bytes]] = []
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            names = set(zf.namelist())
            if "word/document.xml" not in names or "word/_rels/document.xml.rels" not in names:
                return out
            rel_root = etree.fromstring(zf.read("word/_rels/document.xml.rels"))
            rel_map = {}
            for rel in rel_root.findall(DOC_REL_XPATH):
                rid = rel.get("Id")
                target = rel.get("Target", "")
                if rid and target:
                    if not target.startswith("word/"):
                        target = "word/" + target.lstrip("/")
                    rel_map[rid] = target
            doc_root = etree.fromstring(zf.read("word/document.xml"))
            idx = 0
            for blip in doc_root.findall(BLIP_XPATH):
                rid = blip.get(f"{{{R_NS}}}embed")
                if not rid or rid not in rel_map:
                    continue
                target = rel_map[rid]
                if target not in names:
                    continue
                out.append((f"image_{idx:03d}_{Path(target).name}", zf.read(target)))
                idx += 1
    except zipfile.BadZipFile:
        logging.warning(f"跳过无效的压缩包结构 (可能是0字节空文件): {docx_path}")
    except Exception as e:
        logging.warning(f"提取图片时发生未知错误: {docx_path} | 错误信息: {e}")
    return out


def get_ocr_engine():
    try:
        from rapidocr_onnxruntime import RapidOCR
    except Exception as e:
        raise ImportError(
            "未安装 rapidocr_onnxruntime。请先执行: pip install rapidocr-onnxruntime onnxruntime opencv-python-headless") from e
    return RapidOCR()


def ocr_image_bytes(image_bytes: bytes, engine=None) -> str:
    if engine is None:
        engine = get_ocr_engine()
    image = Image.open(BytesIO(image_bytes)).convert("RGB")
    arr = np.array(image)
    result = engine(arr)
    if isinstance(result, tuple) and len(result) >= 1:
        result = result[0]
    lines = []
    if isinstance(result, list):
        for item in result:
            if isinstance(item, (list, tuple)) and len(item) >= 2 and isinstance(item[1], str):
                lines.append(item[1])
    return normalize_spaces("\n".join(lines))


def make_context_snippet(text: str, start: int = 0, end: Optional[int] = None, width: int = 80) -> str:
    t = normalize_spaces(text)
    if not t:
        return ""
    if end is None:
        end = start
    lo = max(0, start - width // 2)
    hi = min(len(t), end + width // 2)
    return t[lo:hi]


# ===================================================================
# 核心修复 1：兼容标签字与字之间出现空格的排版（例如："出 生 地 点"）
# ===================================================================
def label_regex(labels: Iterable[str]) -> re.Pattern:
    spaced_labels = []
    for x in sorted(set(labels), key=len, reverse=True):
        spaced = r"\s*".join(re.escape(char) for char in x)
        spaced_labels.append(spaced)
    escaped = "|".join(spaced_labels)
    return re.compile(rf"(?P<label>{escaped})\s*[：:]\s*", re.IGNORECASE)


# ===================================================================
# 核心修复 2：将带空格的标签还原为标准标签，保障字典提取准确度
# ===================================================================
def extract_kv_segments(text: str, labels: Iterable[str]) -> List[Tuple[str, str, int, int]]:
    text = normalize_spaces(text)
    if not text:
        return []

    # 1. 扫描出所有已知标签（含入院时间、性别等），充当“全局刹车点”
    all_pattern = label_regex(ALL_KNOWN_LABELS)
    all_matches = list(all_pattern.finditer(text))

    # 2. 扫描我们需要真正提取的目标标签（出生地点、住址等）
    target_pattern = label_regex(labels)
    target_matches = list(target_pattern.finditer(text))

    segments = []
    for m in target_matches:
        label = re.sub(r"\s+", "", m.group("label"))
        start_value = m.end()

        # 3. 核心修复：往下找，只要遇到“任何一个”已知标签，就立刻停下！
        end_value = len(text)
        for am in all_matches:
            if am.start() >= start_value:
                end_value = am.start()
                break

        # 截取纯净的值
        raw_value = text[start_value:end_value].strip(" \t\r\n，,；;。|")
        # 清除单元格内手抖敲出的多余换行，压平成干净的一段话
        raw_value = re.sub(r"\s+", " ", raw_value).strip()

        if raw_value:
            segments.append((label, raw_value, m.start(), end_value))

    return segments


def build_patient_rules_from_hits(hits_rows: List[dict], patient_registry_rows: List[dict]) -> Dict[str, dict]:
    rules = {}
    for row in patient_registry_rows:
        rules[row["patient_name"]] = {
            "patient_id": row["patient_id"],
            "replace_patient_terms": {row["patient_name"]},
            "delete_values": set(),
        }
    for row in hits_rows:
        p_name = row.get("patient_name", "")
        if p_name not in rules:
            continue
        raw = row.get("raw_value", "").strip()
        action = row.get("action_suggestion", "")
        if action == "replace_with_patient_id" and raw:
            rules[p_name]["replace_patient_terms"].add(raw)
        elif action == "delete_field_value" and raw:
            rules[p_name]["delete_values"].add(raw)
    return rules


def load_registry_map(path: Path, key_field: str, id_field: str) -> Dict[str, str]:
    rows = safe_read_csv(path)
    out = {}
    for row in rows:
        if not truthy(row.get("confirmed", "1")):
            continue
        key = normalize_spaces(row.get(key_field, ""))
        val = normalize_spaces(row.get(id_field, ""))
        if key and val:
            out[key] = val
    return out


def sorted_replacements(mapping: Dict[str, str]) -> List[Tuple[str, str]]:
    return sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)


def compact_text_preserving_basic(text: str, src: str, dst: str) -> str:
    chars = list(src)
    if len(chars) < 2:
        return text
    spaced = r"\s*".join(map(re.escape, chars))
    return re.sub(spaced, dst, text)


def replace_terms(text: str, replacements: List[Tuple[str, str]]) -> str:
    out = text
    for src, dst in replacements:
        if not src:
            continue
        out = re.sub(re.escape(src), dst, out)
        compact_src = compact_text(src)
        if compact_src and compact_src != src:
            out = compact_text_preserving_basic(out, src, dst)
    return out


def remove_label_value_pairs(text: str, labels: Iterable[str]) -> str:
    labels = sorted(set(labels), key=len, reverse=True)
    if not labels:
        return text

    # 构建包含可选空格的匹配标签
    spaced_labels = [r"\s*".join(re.escape(char) for char in x) for x in labels]
    escaped = "|".join(spaced_labels)

    # 构建包含可选空格的先行断言标签
    spaced_lookaheads = [r"\s*".join(re.escape(char) for char in x) for x in ALL_KNOWN_LABELS]
    lookahead_labels = "|".join(spaced_lookaheads)

    text = normalize_spaces_keep_tabs(text)
    if not text:
        return text

    # 安全策略：只在单行内删除 label:value，避免把后续临床段落/嵌套表格误删。
    line_pattern = re.compile(
        rf"(?P<label>{escaped})\s*[：:]\s*(?P<value>.*?)(?=(?:{lookahead_labels})\s*[：:]|$)",
        re.IGNORECASE,
    )
    out_lines: List[str] = []
    for line in text.split("\n"):
        prev = None
        cur = line
        while prev != cur:
            prev = cur
            cur = line_pattern.sub("", cur)
            cur = re.sub(r"[ \t]{2,}", " ", cur)
            cur = re.sub(r"(^|[ ])([，,；;|]+)", r"\1", cur)
            cur = re.sub(r"([，,；;|]+)([ ]|$)", r"\2", cur)
            cur = cur.strip(" \t|，,；;")
        out_lines.append(cur)

    out = "\n".join(out_lines)
    out = re.sub(r"\n{3,}", "\n\n", out).strip(" \n\t|，,；;")
    return out


def cleanup_text(text: str) -> str:
    text = normalize_spaces_keep_tabs(text)
    text = re.sub(r"[|]{2,}", "|", text)
    text = re.sub(r"[，,；;]{2,}", "，", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def deidentify_text(text: str, patient_name: str, patient_id: str, patient_rules: dict, doctor_map: Dict[str, str],
                    hospital_map: Dict[str, str]) -> str:
    text = normalize_spaces_keep_tabs(text)
    if not text:
        return ""
    text = replace_terms(text, sorted_replacements(hospital_map))
    text = replace_terms(text, sorted_replacements(doctor_map))
    pat_terms = sorted(patient_rules.get("replace_patient_terms", {patient_name}), key=len, reverse=True)
    text = replace_terms(text, [(term, patient_id) for term in pat_terms])
    text = remove_label_value_pairs(text, REMOVE_VALUE_LABELS)
    for raw in sorted(patient_rules.get("delete_values", set()), key=len, reverse=True):
        raw = normalize_spaces(raw)
        if not raw:
            continue
        high_risk = bool(re.search(r"\d{4,}", raw)) or len(raw) >= 6 or any(
            k in raw for k in ["省", "市", "区", "县", "路", "街", "村", "镇", "号", "院", "室"])
        if high_risk:
            text = re.sub(re.escape(raw), "", text)
    return cleanup_text(text)


def dump_json(path: Path, obj) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
